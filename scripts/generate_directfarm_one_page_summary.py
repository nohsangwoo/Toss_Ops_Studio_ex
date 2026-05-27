from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "docs" / "directfarm-one-page-summary.docx"

FONT = "Malgun Gothic"
BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F3F4F6"
BORDER = "D9D9D9"


def set_font(run, size=9.5, bold=False, color=DARK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_p(p, before=0, after=3, line=1.05):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def para(doc, text, size=9.5, bold=False, color=DARK, after=3, align=None):
    p = doc.add_paragraph()
    set_p(p, after=after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def rule(doc):
    p = doc.add_paragraph()
    set_p(p, after=5)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    p_pr.append(border)


def borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def shade(cell, fill=LIGHT_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=8.8, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    set_p(p, after=0, line=1.0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)


def cell_width(cell, cm):
    cell.width = Cm(cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for idx, h in enumerate(headers):
        cell_text(t.rows[0].cells[idx], h, bold=True, fill=LIGHT_FILL)
        cell_width(t.rows[0].cells[idx], widths[idx])
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)
            cell_width(cells[idx], widths[idx])
    para(doc, "", after=1)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_p(p, after=1, line=1.0)
    run = p.add_run(text)
    set_font(run, size=9)


def heading(doc, text):
    para(doc, text, size=11, bold=True, color=BLACK, after=3)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    para(doc, "DirectFarm MVP 한 장 요약", size=19, bold=True, color=BLACK, after=2)
    para(doc, "야외 무인 주문 키오스크 및 도매처 직배송 O2O 웹 플랫폼", size=10.5, color=MUTED, after=4)
    rule(doc)

    table(
        doc,
        ["항목", "내용"],
        [
            ["제안 금액", "5,440,000원 / VAT 별도"],
            ["작업 기간", "40일"],
            ["투입 인력", "PM 1명, PL 1명"],
            ["수행사", "주식회사 럿지 (LUDGI Inc.)"],
        ],
        [3.0, 14.5],
    )

    heading(doc, "프로젝트 목적")
    para(
        doc,
        "태블릿 키오스크에서 상품 선택, 배송지 입력, QR 결제, 주문 저장, 도매처 전달, 관리자 확인까지 한 흐름으로 검증하는 MVP 구축.",
        size=9.3,
        after=5,
    )

    heading(doc, "핵심 구현 범위")
    for item in [
        "키오스크 상품 리스트, 상품 선택 팝업, 수량 선택, 배송지 입력",
        "Toss Payments QR/간편결제, 서버 금액 검증, 결제 성공/실패 처리",
        "상품별 도매처 매핑, 결제 완료 주문 전송 로그 저장",
        "관리자 상품/도매처/주문/결제/전송 상태 관리, CSV 다운로드",
        "Vercel 배포, PostgreSQL 기반 DB, 운영 문서 및 README 제공",
    ]:
        bullet(doc, item)

    heading(doc, "제안 스택")
    table(
        doc,
        ["영역", "기술"],
        [
            ["Frontend", "Next.js 16, React 19, TypeScript, Tailwind CSS, shadcn/ui"],
            ["Backend/DB", "Next.js API Routes, Server Actions, Prisma ORM, Neon PostgreSQL"],
            ["Payment/Security", "Toss Payments, Cloudflare Turnstile, NextAuth Admin Role"],
            ["Deploy/Expansion", "Vercel, Supabase 대체 가능, Firebase Auth/FCM 확장 가능"],
        ],
        [3.5, 14.0],
    )

    heading(doc, "수행 강점")
    for item in [
        "PG 심사, 결제 승인, 취소/환불, 웹훅, 관리자 결제 관리 경험 다수",
        "PayPal, Toss Payments, KG이니시스, 나이스페이, 쿠키페이 등 PG 연동 경험",
        "카카오 알림톡, SMS, SendGrid 등 외부 발송 API 연동 경험",
        "태블릿 우선 UX로 시작하되 모바일/PC/앱 확장까지 고려한 구조 제안",
    ]:
        bullet(doc, item)

    heading(doc, "확인 링크")
    table(
        doc,
        ["구분", "URL"],
        [
            ["키오스크 데모", "https://tossops.bsclinic.xyz/directfarm"],
            ["관리자 데모", "https://tossops.bsclinic.xyz/directfarm/admin"],
            ["제안 리포트", "https://tossops.bsclinic.xyz/directfarm/proposal"],
            ["전용 README", "https://github.com/nohsangwoo/Toss_Ops_Studio_ex/tree/main/docs/projects/directfarm"],
            ["미팅 제안서", "https://tossops.bsclinic.xyz/docs/directfarm-meeting-proposal-company.docx"],
        ],
        [3.2, 14.3],
    )

    heading(doc, "미팅에서 바로 확정할 항목")
    para(
        doc,
        "PG 결제수단 범위 / 알림톡·SMS provider / 도매처 전달 항목 / 키오스크 기기 환경 / 환불·취소 운영 정책",
        size=9.3,
        after=0,
    )

    doc.core_properties.title = "DirectFarm One Page Summary"
    doc.core_properties.author = "LUDGI"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
