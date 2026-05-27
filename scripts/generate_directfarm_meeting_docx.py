from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "docs" / "directfarm-meeting-proposal.docx"
ASSET_DIR = ROOT / "public" / "images" / "directfarm" / "proposal"

FONT = "Malgun Gothic"
BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F3F4F6"
BORDER = "D9D9D9"
TOTAL_WIDTH_CM = 17.0


def set_east_asia_font(run, font_name: str = FONT) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run(run, size: float = 10.5, bold: bool = False, color: RGBColor = DARK) -> None:
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph(p, before: float = 0, after: float = 5, line_spacing: float = 1.12) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing


def add_para(doc: Document, text: str = "", size: float = 10.5, bold: bool = False,
             color: RGBColor = DARK, after: float = 5, before: float = 0,
             align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(p, before=14 if level == 1 else 10, after=6)
    run = p.add_run(text)
    set_run(run, size=15 if level == 1 else 12.5, bold=True, color=BLACK)


def add_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=9)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    p_pr.append(border)


def shade_cell(cell, fill: str = LIGHT_FILL) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5,
                  color: RGBColor = DARK, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.text = ""
    set_paragraph(p, after=0, line_spacing=1.12)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, fill=LIGHT_FILL)
        set_cell_width(table.rows[0].cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, after=3, line_spacing=1.12)
        run = p.add_run(item)
        set_run(run, size=10.2)


def add_link(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=3)
    label_run = p.add_run(f"{label}: ")
    set_run(label_run, size=10, bold=True, color=BLACK)
    url_run = p.add_run(url)
    set_run(url_run, size=9.5, color=MUTED)


def add_cover_page(doc: Document) -> None:
    add_para(doc, "주식회사 럿지", size=11, bold=True, color=MUTED, after=4)
    add_para(doc, "DIRECTFARM", size=28, bold=True, color=BLACK, after=2)
    add_para(doc, "MVP 개발 제안 브리프", size=24, bold=True, color=BLACK, after=12)
    add_rule(doc)
    add_para(
        doc,
        "야외 무인 주문(키오스크형) 및 도매처 직배송 O2O 웹 플랫폼 MVP 개발",
        size=14,
        color=DARK,
        after=16,
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["제출사", "주식회사 럿지 (LUDGI Inc.)"],
            ["문서 목적", "미팅 전 프로젝트 이해, 수행 범위, 견적 및 구현 방향 공유"],
            ["제안 금액", "5,440,000원 / VAT 별도"],
            ["제안 기간", "40일"],
            ["투입 인력", "PM 1명, PL 1명"],
            ["작성일", "2026.05.27"],
        ],
        [3.5, 13.0],
    )
    add_para(
        doc,
        "본 문서는 미팅 검토용으로 작성된 요약 제안서이며, 세부 범위와 일정은 발주처 요구사항 확정 후 조정될 수 있습니다.",
        size=9.5,
        color=MUTED,
        after=0,
    )
    doc.add_page_break()


def add_company_profile(doc: Document) -> None:
    add_para(doc, "COMPANY PROFILE", size=18, bold=True, color=BLACK, after=2)
    add_para(doc, "주식회사 럿지 (LUDGI Inc.)", size=13, bold=True, color=DARK, after=8)
    add_para(
        doc,
        "주식회사 럿지는 웹/앱 제품 개발, 공공·민간 SI, 스타트업 MVP, AI·데이터 솔루션, 클라우드 인프라 구축을 수행하는 소프트웨어 개발 파트너입니다.",
    )
    add_bullets(
        doc,
        [
            "공공기관 SI 및 민간 프로젝트 수행 경험을 기반으로 기획부터 배포까지 일괄 대응",
            "Next.js, React, TypeScript, Node.js, Python, Flutter 기반의 풀스택 제품 개발",
            "AWS, GCP, Firebase, Docker, Kubernetes, Vercel 등 운영 인프라 경험",
            "LLM, RAG, Vector Search, Computer Vision 등 제품형 AI 기능 연동 가능",
            "스타트업 MVP와 운영형 관리자 시스템을 모두 고려한 실무 중심 개발 방식",
        ],
    )
    add_table(
        doc,
        ["항목", "회사 정보"],
        [
            ["법인명", "LUDGI Inc. (주식회사 럿지)"],
            ["대표자", "노상우"],
            ["설립", "2024년"],
            ["사업자등록번호", "307-88-03283"],
            ["DUNS Number", "963415644"],
            ["주소", "인천광역시 연수구 인천타워대로 323, 에이동 20층"],
            ["연락처", "010-3006-9310"],
            ["이메일", "milli@molluhub.com"],
        ],
        [3.5, 13.0],
    )
    add_heading(doc, "회사 역량 요약", 2)
    add_table(
        doc,
        ["영역", "주요 역량"],
        [
            ["Full-stack Engineering", "React, Next.js, Flutter, Node.js, Python 기반 웹/앱 제품 개발"],
            ["Cloud & DevOps", "AWS, GCP, Firebase, Docker, Kubernetes, Vercel 기반 배포·운영"],
            ["Public Sector SI", "나라장터, KEPCO, KDN 등 공공/규정 대응 경험"],
            ["Startup MVP", "기획, 디자인, 개발, 출시까지 MVP 구축 일괄 수행"],
            ["Technical Consulting", "아키텍처 설계, 기술 스택 선정, 코드 리뷰, 인프라 점검"],
        ],
        [4.7, 11.8],
    )
    add_heading(doc, "본 프로젝트와의 적합성", 2)
    add_bullets(
        doc,
        [
            "결제, 주문, 관리자, 배포를 한 팀에서 연결해 MVP 속도를 높일 수 있습니다.",
            "PG/알림/DB/관리자 권한 등 실제 운영에서 자주 문제가 되는 지점을 사전에 구조화합니다.",
            "태블릿 키오스크에서 시작하되, 향후 모바일 앱과 PC 운영 화면까지 확장 가능한 구조를 제안합니다.",
        ],
    )
    add_link(doc, "회사 소개", "https://info.ludgi.ai/company")
    add_link(doc, "회사 홈페이지", "https://info.ludgi.ai/")
    doc.add_page_break()


def convert_for_docx(src: Path, tmp_dir: Path, name: str, max_width: int = 1500) -> Path:
    image = Image.open(src)
    image = ImageOps.exif_transpose(image)
    if image.width > max_width:
        ratio = max_width / image.width
        image = image.resize((max_width, int(image.height * ratio)))
    image = image.convert("RGB")
    out = tmp_dir / f"{name}.png"
    image.save(out, "PNG", optimize=True)
    return out


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(cap, after=8)
    r = cap.add_run(caption)
    set_run(r, size=8.8, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.color.rgb = BLACK

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DirectFarm MVP Meeting Brief")
    set_run(run, size=8, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Confidential meeting material | 2026.05.27")
    set_run(run, size=8, color=MUTED)


def build_doc() -> None:
    doc = Document()
    configure_document(doc)

    doc.core_properties.title = "DirectFarm 미팅 제안서"
    doc.core_properties.subject = "야외 무인 주문 및 도매처 직배송 O2O 웹 플랫폼 MVP 개발"
    doc.core_properties.author = "LUDGI"

    add_cover_page(doc)
    add_company_profile(doc)

    add_para(doc, "DIRECTFARM MVP 제안 브리프", size=22, bold=True, color=BLACK, after=4)
    add_para(
        doc,
        "야외 무인 주문(키오스크형) 및 도매처 직배송 O2O 웹 플랫폼 MVP 개발",
        size=13,
        color=MUTED,
        after=10,
    )
    add_rule(doc)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["제안 금액", "5,440,000원 / VAT 별도"],
            ["제안 기간", "40일"],
            ["투입 인력", "PM 1명, PL 1명"],
            ["수행 역할", "PM 및 제작총괄"],
            ["목업 제작 기록", "2026년 5월부터 약 1개월간 구현"],
            ["핵심 산출물", "키오스크 주문 화면, QR 결제, 도매처 전송 로그, 관리자 운영 화면, 배포 문서"],
        ],
        [3.5, 13.0],
    )

    add_heading(doc, "1. 제안 요약", 1)
    add_para(
        doc,
        "본 제안은 야외 태블릿 키오스크에서 고객이 상품을 선택하고 QR 간편결제로 결제하면, 주문 정보가 도매처로 전달되는 산지직송 O2O MVP 구축을 목표로 합니다.",
    )
    add_bullets(
        doc,
        [
            "단건 빠른 주문 흐름에 집중하여 초기 사업성 검증 비용과 시간을 줄입니다.",
            "결제 금액은 서버에서 재검증하고, 결제 결과는 DB에 저장해 운영 안정성을 확보합니다.",
            "상품, 가격, 도매처, 주문 상태, 전송 상태는 관리자 화면에서 직접 관리합니다.",
            "태블릿 우선 UX로 설계하되, 모바일/PC 및 향후 앱 확장 여지를 남깁니다.",
        ],
    )

    add_heading(doc, "2. 수행사 역량 요약", 1)
    add_para(
        doc,
        "법인 개발사로서 결제 연동, 주문/상품 관리자, 외부 API 연동, 배포 운영까지 한 흐름으로 수행 가능한 체계를 보유하고 있습니다.",
    )
    add_bullets(
        doc,
        [
            "PayPal, Toss Payments, KG이니시스, 나이스페이, 쿠키페이 등 국내외 PG 연동 경험",
            "PG사 심사 대응, 결제 승인, 취소/환불, 웹훅, 관리자 결제 관리 기능 구현 경험",
            "네트워크가 불안정한 환경을 고려한 웹훅 기반 결제 상태 보정 및 로그 관리 경험",
            "카카오 알림톡, SMS, SendGrid 등 메시징/메일 발송 API 연동 경험",
            "통합 쇼핑몰 앱, 주문 관리자, 상품 관리자, 결제 운영 화면 구축 경험",
        ],
    )

    add_heading(doc, "3. 프로젝트 이해", 1)
    add_table(
        doc,
        ["구분", "이해 내용"],
        [
            ["문제", "야외 무인 판매 환경에서는 직원 상주가 어렵고, 결제 이후 주문 정보를 도매처에 별도로 전달해야 하므로 운영 누락 가능성이 있습니다."],
            ["목표", "태블릿 한 대에서 상품 선택, 배송지 입력, QR 결제, 주문 저장, 도매처 전달, 관리자 확인까지 이어지는 MVP를 구축합니다."],
            ["주안점", "복잡한 장바구니보다 단건 빠른 결제와 관리자 운영 편의성을 우선하고, 실운영 확장 시 알림톡/SMS와 앱 확장을 고려합니다."],
        ],
        [3.0, 13.5],
    )

    add_heading(doc, "4. 수행 범위", 1)
    add_table(
        doc,
        ["영역", "구현 범위"],
        [
            ["키오스크 화면", "상품 리스트, 사진 중심 카드, 상품 상세 팝업, 수량 선택, 배송지 입력, 결제 진행"],
            ["결제/보안", "Toss Payments QR 결제위젯, 서버 금액 검증, 결제 성공/실패 처리, Cloudflare Turnstile 인증"],
            ["도매처 전달", "상품별 도매처 매핑, 결제 완료 주문 payload 생성, 전송 로그 저장, 추후 Solapi/Aligo 전환 구조"],
            ["관리자 화면", "상품 등록/수정, 판매가/도매가 관리, 도매처 관리, 주문/결제/전송 상태 조회, CSV 다운로드"],
            ["배포/문서", "Vercel 배포, PostgreSQL 연동, 운영 환경변수 정리, README 및 제안/견적 리포트 제공"],
        ],
        [3.2, 13.3],
    )

    add_heading(doc, "5. 기술 구성", 1)
    add_table(
        doc,
        ["구분", "제안 스택"],
        [
            ["Frontend", "Next.js 16 App Router, React 19, TypeScript, Tailwind CSS, shadcn/ui"],
            ["Backend", "Next.js Route Handlers, Server Actions, Prisma ORM"],
            ["Database", "Neon PostgreSQL 기본 제안, 발주처 상황에 따라 Supabase PostgreSQL 대체 가능"],
            ["Payment", "Toss Payments 결제위젯, QR 간편결제, 승인/실패/영수증 처리"],
            ["Security", "Cloudflare Turnstile, 관리자 권한 관리, 서버 결제 금액 검증"],
            ["Deploy", "Vercel, 커스텀 도메인, 환경변수 기반 운영 설정"],
            ["Expansion", "회원/앱 확장이 필요한 경우 Firebase Auth, FCM 등을 선택 적용 가능"],
        ],
        [3.2, 13.3],
    )

    add_heading(doc, "6. 일정 계획", 1)
    add_table(
        doc,
        ["단계", "기간", "주요 내용"],
        [
            ["1. 기획 및 요구사항 정리", "1주차", "주문 흐름, 관리자 범위, PG/알림 정책, 데이터 항목 확정"],
            ["2. UI/UX 설계", "1-2주차", "태블릿 키오스크 화면, 상품 팝업, 관리자 화면 구조 설계"],
            ["3. 주문·결제 개발", "2-3주차", "상품 선택, 배송지 입력, QR 결제, 서버 승인/실패 처리 구현"],
            ["4. 관리자·전송 기능", "3-5주차", "상품/도매처/주문 관리, 전송 로그, CSV 다운로드 구현"],
            ["5. 검수 및 배포", "5-6주차", "운영 환경 배포, 결제/주문/관리자 테스트, 문서 정리"],
        ],
        [4.2, 2.4, 9.9],
    )

    add_heading(doc, "7. 비용 및 조건", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["개발 비용", "5,440,000원 / VAT 별도"],
            ["외부 비용", "PG 수수료, 문자/알림톡 발송비, Vercel/DB 등 외부 서비스 이용료는 별도"],
            ["제외 범위", "도매처별 자동 정산, 복수 상품 장바구니, 실물 카드 단말기 연동은 별도 협의"],
            ["운영 전환", "실제 알림톡/SMS 발송은 Solapi 또는 Aligo 계정과 발송 키 발급 후 연결"],
        ],
        [3.2, 13.3],
    )

    add_heading(doc, "8. 참고 포트폴리오 및 링크", 1)
    add_link(doc, "DirectFarm 키오스크 데모", "https://tossops.bsclinic.xyz/directfarm")
    add_link(doc, "DirectFarm 관리자 데모", "https://tossops.bsclinic.xyz/directfarm/admin")
    add_link(doc, "DirectFarm 제안/견적 리포트", "https://tossops.bsclinic.xyz/directfarm/proposal")
    add_link(doc, "DirectFarm 전용 README", "https://github.com/nohsangwoo/Toss_Ops_Studio_ex/tree/main/docs/projects/directfarm")
    add_link(doc, "Toss Payments 결제 관리 포트폴리오", "https://tossops.bsclinic.xyz/showcase")
    add_link(doc, "GitHub 저장소", "https://github.com/nohsangwoo/Toss_Ops_Studio_ex")

    doc.add_page_break()
    add_heading(doc, "9. 화면 증빙", 1)
    add_para(
        doc,
        "아래 화면은 미팅 전 검토용 목업에서 캡처한 주요 흐름입니다. 문서 톤을 맞추기 위해 흑백 이미지로 정리했습니다.",
        color=MUTED,
    )

    with TemporaryDirectory() as temp:
        temp_dir = Path(temp)
        figures = [
            ("kiosk-flow.webp", "키오스크 홈, 상품 카탈로그, 주문 흐름 전체"),
            ("product-detail.webp", "상품 선택 후 수량 및 결제 전 확인"),
            ("checkout-before-payment.webp", "배송지 입력, Turnstile 인증, Toss 결제위젯"),
            ("qr-payment.webp", "Toss Payments QR 간편결제 호출"),
            ("admin-flow.webp", "관리자 상품/도매처/주문 운영 화면"),
            ("sales-receipt.webp", "Toss Payments 매출전표 및 결제 증빙"),
        ]
        for idx, (filename, caption) in enumerate(figures, start=1):
            converted = convert_for_docx(ASSET_DIR / filename, temp_dir, f"figure-{idx}")
            add_figure(doc, converted, f"Figure {idx}. {caption}")
            if idx in (2, 4):
                doc.add_page_break()

    add_heading(doc, "10. 미팅 시 확인하면 좋은 사항", 1)
    add_bullets(
        doc,
        [
            "실제 사용할 PG사와 결제수단 범위: QR 간편결제, 카드, 가상계좌 등",
            "알림톡/SMS 발송 provider: Solapi, Aligo, 기존 계약 서비스 여부",
            "도매처에 전달할 주문 정보 항목과 개인정보 보관 기간",
            "태블릿 기기 해상도, 브라우저/키오스크 모드 운영 방식",
            "상품 등록/가격 수정 권한, 관리자 계정 수, 운영자별 권한 분리 필요 여부",
            "결제 취소/환불 정책과 고객 문의 대응 흐름",
            "초기 MVP 이후 앱, 회원가입, 푸시 알림, 재구매 기능 확장 계획",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
